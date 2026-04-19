"""
EduForge — Ingestion Service
Handles: file parsing (PDF, DOCX, TXT, MD) → text extraction
         → semantic chunking → embedding → ChromaDB upsert
Uses ChromaDB in persistent local mode (no Docker needed).
"""
from __future__ import annotations

import re
import time
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

import chromadb
from chromadb.config import Settings as ChromaSettings
from sentence_transformers import SentenceTransformer

from core.config import settings
from core.utils import get_logger

logger = get_logger("eduforge.ingestion")


# ── Text Extraction ───────────────────────────────────────────────────────────

class TextExtractor:
    @staticmethod
    def extract(file_path: Path, file_type: Optional[str] = None) -> str:
        suffix = (file_type or file_path.suffix).lower().lstrip(".")
        extractors = {
            "pdf":  TextExtractor._extract_pdf,
            "docx": TextExtractor._extract_docx,
            "txt":  TextExtractor._extract_txt,
            "md":   TextExtractor._extract_txt,
        }
        fn = extractors.get(suffix)
        if fn is None:
            raise ValueError(f"Unsupported file type: {suffix}")
        logger.info("Extracting text", extra={"file": str(file_path), "type": suffix})
        return fn(file_path)

    @staticmethod
    def extract_txt(file_path: Path) -> str:
        return TextExtractor._extract_txt(file_path)

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i+1}]\n{text}")
            return "\n\n".join(pages)
        except ImportError:
            raise RuntimeError("pypdf not installed — run: pip install pypdf")

    @staticmethod
    def _extract_docx(path: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)
        except ImportError:
            raise RuntimeError("python-docx not installed")

    @staticmethod
    def _extract_txt(path: Path) -> str:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return path.read_text(encoding=enc)
            except (UnicodeDecodeError, FileNotFoundError):
                continue
        raise RuntimeError(f"Cannot decode file: {path}")


# ── Chunking ──────────────────────────────────────────────────────────────────

class TextChunker:
    def __init__(
        self,
        chunk_size: int = settings.CHUNK_SIZE,
        overlap: int = settings.CHUNK_OVERLAP,
    ):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, text: str) -> List[Dict[str, Any]]:
        text = self._clean(text)
        paragraphs = self._split_paragraphs(text)
        chunks: List[Dict[str, Any]] = []
        current: List[str] = []
        current_len = 0
        idx = 0

        for para in paragraphs:
            para_tokens = self._token_count(para)
            if current_len + para_tokens > self.chunk_size and current:
                chunk_text = " ".join(current).strip()
                if chunk_text:
                    chunks.append({
                        "content": chunk_text,
                        "chunk_index": idx,
                        "token_count": current_len,
                        "metadata": {"chunk_index": idx},
                    })
                    idx += 1
                keep = self._keep_overlap(current)
                current = keep
                current_len = sum(self._token_count(s) for s in current)

            current.append(para)
            current_len += para_tokens

        if current:
            chunk_text = " ".join(current).strip()
            if chunk_text:
                chunks.append({
                    "content": chunk_text,
                    "chunk_index": idx,
                    "token_count": current_len,
                    "metadata": {"chunk_index": idx},
                })

        logger.info("Chunked document", extra={"num_chunks": len(chunks)})
        return chunks

    @staticmethod
    def _clean(text: str) -> str:
        text = re.sub(r"\s*\n\s*\n\s*", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    @staticmethod
    def _split_paragraphs(text: str) -> List[str]:
        paragraphs = re.split(r"\n{2,}", text)
        return [p.replace("\n", " ").strip() for p in paragraphs if p.strip()]

    @staticmethod
    def _token_count(text: str) -> int:
        return max(1, len(text) // 4)

    def _keep_overlap(self, sentences: List[str]) -> List[str]:
        kept = []
        total = 0
        for s in reversed(sentences):
            t = self._token_count(s)
            if total + t > self.overlap:
                break
            kept.insert(0, s)
            total += t
        return kept


# ── Embedding ─────────────────────────────────────────────────────────────────

class EmbeddingEngine:
    _instance: Optional["EmbeddingEngine"] = None

    def __init__(self, model_name: str = settings.EMBEDDING_MODEL):
        logger.info("Loading embedding model", extra={"model": model_name})
        t0 = time.time()
        self.model = SentenceTransformer(
            model_name,
            cache_folder=str(settings.MODEL_CACHE_DIR),
        )
        self.model_name = model_name
        self.dim = self.model.get_sentence_embedding_dimension()
        logger.info("Embedding model loaded", extra={"elapsed_s": round(time.time()-t0, 2)})

    @classmethod
    def get_instance(cls) -> "EmbeddingEngine":
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def embed(self, texts: List[str], batch_size: int = 32) -> List[List[float]]:
        if not texts:
            return []
        vecs = self.model.encode(
            texts,
            batch_size=batch_size,
            show_progress_bar=False,
            normalize_embeddings=True,
        )
        return vecs.tolist()

    def embed_one(self, text: str) -> List[float]:
        return self.embed([text])[0]


# ── Vector Store ──────────────────────────────────────────────────────────────

class VectorStore:
    """
    ChromaDB in persistent local mode — no Docker/HTTP needed.
    Data stored in data/processed/chroma/
    """
    _instance: Optional["VectorStore"] = None

    def __init__(self):
        chroma_dir = settings.PROCESSED_DIR / "chroma"
        chroma_dir.mkdir(parents=True, exist_ok=True)
        logger.info("Initialising local ChromaDB", extra={"path": str(chroma_dir)})
        self.client = chromadb.PersistentClient(path=str(chroma_dir))
        self.collection = self.client.get_or_create_collection(
            name=settings.CHROMA_COLLECTION_NAME,
            metadata={"hnsw:space": "cosine"},
        )
        logger.info("VectorStore ready", extra={"collection": settings.CHROMA_COLLECTION_NAME})

    @classmethod
    def get_instance(cls) -> "VectorStore":
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def upsert(
        self,
        chunk_ids: List[str],
        embeddings: List[List[float]],
        documents: List[str],
        metadatas: List[Dict[str, Any]],
    ) -> None:
        self.collection.upsert(
            ids=chunk_ids,
            embeddings=embeddings,
            documents=documents,
            metadatas=metadatas,
        )

    def query(
        self,
        query_embedding: List[float],
        n_results: int = 5,
        where: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        kwargs: Dict[str, Any] = {
            "query_embeddings": [query_embedding],
            "n_results": min(n_results, max(1, self.collection.count())),
            "include": ["documents", "metadatas", "distances"],
        }
        if where:
            kwargs["where"] = where

        result = self.collection.query(**kwargs)
        hits = []
        for doc, meta, dist in zip(
            result["documents"][0],
            result["metadatas"][0],
            result["distances"][0],
        ):
            hits.append({
                "content": doc,
                "metadata": meta,
                "score": 1.0 - dist,
            })
        return hits

    def delete_by_material(self, material_id: str) -> None:
        try:
            self.collection.delete(where={"material_id": material_id})
        except Exception:
            pass

    def count(self) -> int:
        return self.collection.count()

    def health(self) -> bool:
        try:
            self.collection.count()
            return True
        except Exception:
            return False


# ── Ingestion Pipeline ────────────────────────────────────────────────────────

class IngestionPipeline:
    def __init__(
        self,
        extractor: Optional[TextExtractor] = None,
        chunker: Optional[TextChunker] = None,
        embedder: Optional[EmbeddingEngine] = None,
        vector_store: Optional[VectorStore] = None,
    ):
        self.extractor    = extractor    or TextExtractor()
        self.chunker      = chunker      or TextChunker()
        self.embedder     = embedder     or EmbeddingEngine.get_instance()
        self.vector_store = vector_store or VectorStore.get_instance()

    def run(
        self,
        file_path: Path,
        material_id: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        logger.info("Ingestion started", extra={"material_id": material_id})
        t0 = time.time()

        raw_text = self.extractor.extract(file_path)
        if not raw_text.strip():
            raise ValueError("No text extracted from document")

        chunks = self.chunker.chunk(raw_text)
        if not chunks:
            raise ValueError("No chunks produced from document")

        texts = [c["content"] for c in chunks]
        embeddings = self.embedder.embed(texts)

        chunk_ids = [str(uuid.uuid4()) for _ in chunks]
        meta_base = {"material_id": material_id, **(metadata or {})}
        metadatas = [
            {**meta_base, "chunk_index": c["chunk_index"], "token_count": c["token_count"]}
            for c in chunks
        ]

        self.vector_store.upsert(chunk_ids, embeddings, texts, metadatas)

        for chunk, cid in zip(chunks, chunk_ids):
            chunk["vector_id"] = cid
            chunk["material_id"] = material_id

        elapsed = round(time.time() - t0, 2)
        logger.info("Ingestion complete", extra={"chunks": len(chunks), "elapsed_s": elapsed})
        return chunks

    def retrieve_context(
        self,
        query: str,
        n_results: int = 5,
        material_ids: Optional[List[str]] = None,
    ) -> List[str]:
        q_emb = self.embedder.embed_one(query)
        where = {"material_id": {"$in": material_ids}} if material_ids else None
        hits = self.vector_store.query(q_emb, n_results=n_results, where=where)
        return [h["content"] for h in hits]